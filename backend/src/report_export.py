from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _clean_filename(source_filename: str | None) -> str:
    stem = Path(source_filename or "").stem.strip()
    return stem or "避坑指南"


def build_report_docx(report_paragraphs: list[str], source_filename: str | None = None) -> bytes:
    document = Document()

    title = document.add_heading("合同避坑指南", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"来源文件：{_clean_filename(source_filename)}")

    for paragraph in report_paragraphs:
        content = paragraph.strip()
        if not content:
            continue

        if content.startswith("## "):
            document.add_heading(content.replace("## ", "", 1).strip(), level=1)
            continue

        if content.startswith("### "):
            document.add_heading(content.replace("### ", "", 1).strip(), level=2)
            continue

        document.add_paragraph(content)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_report_download_name(source_filename: str | None = None) -> str:
    return f"{_clean_filename(source_filename)}_避坑指南.docx"
